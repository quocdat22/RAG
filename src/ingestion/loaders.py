"""
Document loaders for various file types.

This module provides loaders for different document formats (PDF, TXT, CSV, DOCX, XLSX).
Each loader extracts text content and metadata from documents.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

import pandas as pd
from pypdf import PdfReader

from src.core.exceptions import DocumentLoadError, UnsupportedFileTypeError
from src.core.logging import LoggerMixin
from src.core.utils import generate_doc_id, get_file_extension, sanitize_filename


class Document:
    """Represents a loaded document with content and metadata."""

    def __init__(
        self,
        content: str,
        metadata: dict[str, Any],
        doc_id: str | None = None,
    ):
        """
        Initialize document.

        Args:
            content: Document text content
            metadata: Document metadata
            doc_id: Optional document ID (generated if not provided)
        """
        self.content = content
        self.metadata = metadata
        self.doc_id = doc_id or generate_doc_id(
            metadata.get("filename", "unknown"), content
        )

    def __repr__(self) -> str:
        return f"Document(id={self.doc_id}, content_length={len(self.content)})"


class BaseLoader(ABC, LoggerMixin):
    """Base class for document loaders."""

    @abstractmethod
    def load(self, file_path: Path) -> Document:
        """
        Load document from file.

        Args:
            file_path: Path to file

        Returns:
            Loaded document

        Raises:
            DocumentLoadError: If loading fails
        """
        pass

    def _create_metadata(self, file_path: Path) -> dict[str, Any]:
        """
        Create base metadata for document.

        Args:
            file_path: Path to file

        Returns:
            Metadata dictionary
        """
        stat = file_path.stat()

        return {
            "filename": file_path.name,
            "filepath": str(file_path.absolute()),
            "file_type": get_file_extension(file_path.name),
            "file_size_bytes": stat.st_size,
            "created_at": stat.st_ctime,
            "modified_at": stat.st_mtime,
        }


class PDFLoader(BaseLoader):
    """Loader for PDF files."""
    
    def _extract_year_from_metadata(self, reader: PdfReader) -> str:
        """
        Extract publication year from PDF metadata.
        
        Args:
            reader: PDF reader instance
            
        Returns:
            Year as string, or empty string if not found
        """
        import re
        from datetime import datetime
        
        try:
            if not reader.metadata:
                return ""
            
            # Try to extract from creation or modification date
            date_fields = ["/CreationDate", "/ModDate"]
            for field in date_fields:
                date_str = reader.metadata.get(field, "")
                if date_str:
                    # PDF date format: D:YYYYMMDDHHmmSSOHH'mm'
                    year_match = re.search(r"D:(\d{4})", str(date_str))
                    if year_match:
                        return year_match.group(1)
                    
                    # Try parsing as datetime
                    try:
                        if isinstance(date_str, datetime):
                            return str(date_str.year)
                    except:
                        pass
        except Exception as e:
            self.logger.debug(f"Failed to extract year from PDF metadata: {e}")
        
        return ""
    
    def _extract_year_from_content(self, content: str) -> str:
        """
        Extract year from content by looking for patterns.
        
        Args:
            content: PDF text content
            
        Returns:
            Year as string, or empty string if not found
        """
        import re
        
        # Look for common year patterns in first 2000 chars
        sample = content[:2000]
        
        # Pattern: (2024), CVPR 2024, NeurIPS 2023, etc.
        patterns = [
            r"\b(19|20)\d{2}\b",  # Any 4-digit year
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, sample)
            if matches:
                # Return first reasonable year (1990-2030)
                for match in matches:
                    year = int(match) if isinstance(match, str) else int(match)
                    if 1990 <= year <= 2030:
                        return str(year)
        
        return ""
    
    def _format_author(self, author: str) -> str:
        """
        Format author name(s) for citation.
        
        Args:
            author: Author string from PDF metadata
            
        Returns:
            Formatted author string (e.g., "Author et al.")
        """
        if not author:
            return ""
        
        # Split by common separators
        separators = [",", ";", " and ", "&"]
        authors = [author]
        for sep in separators:
            if sep in author:
                authors = author.split(sep)
                break
        
        # Clean whitespace
        authors = [a.strip() for a in authors if a.strip()]
        
        if len(authors) == 0:
            return ""
        elif len(authors) == 1:
            return authors[0]
        elif len(authors) == 2:
            return f"{authors[0]} and {authors[1]}"
        else:
            # Get first author's last name if possible
            first_author = authors[0]
            # Try to extract last name (assumes "First Last" format)
            name_parts = first_author.split()
            if len(name_parts) > 1:
                return f"{name_parts[-1]} et al."
            else:
                return f"{first_author} et al."

    def load(self, file_path: Path) -> Document:
        """Load PDF document."""
        try:
            self.logger.info(f"Loading PDF: {file_path.name}")

            reader = PdfReader(file_path)
            total_pages = len(reader.pages)

            # Extract text from all pages
            text_content = []
            for page_num, page in enumerate(reader.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"[Page {page_num}]\n{page_text}")
                except Exception as e:
                    self.logger.warning(
                        f"Failed to extract text from page {page_num}: {e}"
                    )

            content = "\n\n".join(text_content)

            if not content.strip():
                raise DocumentLoadError(
                    file_path.name, "No text content extracted from PDF"
                )

            # Create metadata
            metadata = self._create_metadata(file_path)
            metadata.update(
                {
                    "total_pages": total_pages,
                    "loader": "PDFLoader",
                }
            )

            # Try to extract PDF metadata
            try:
                if reader.metadata:
                    raw_author = reader.metadata.get("/Author", "")
                    metadata.update(
                        {
                            "title": reader.metadata.get("/Title", ""),
                            "author": raw_author,
                            "author_formatted": self._format_author(raw_author),
                            "subject": reader.metadata.get("/Subject", ""),
                            "creator": reader.metadata.get("/Creator", ""),
                        }
                    )
            except Exception:
                pass
            
            # Extract publication year
            year = self._extract_year_from_metadata(reader)
            if not year:
                # Fallback to content parsing
                year = self._extract_year_from_content(content)
            
            if year:
                metadata["year"] = year

            self.logger.info(
                f"Successfully loaded PDF: {file_path.name} ({total_pages} pages)"
            )

            return Document(content=content, metadata=metadata)

        except Exception as e:
            raise DocumentLoadError(file_path.name, str(e))


class TXTLoader(BaseLoader):
    """Loader for plain text files."""

    def load(self, file_path: Path) -> Document:
        """Load text document."""
        try:
            self.logger.info(f"Loading TXT: {file_path.name}")

            # Try different encodings
            encodings = ["utf-8", "latin-1", "cp1252"]
            content = None

            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                raise DocumentLoadError(
                    file_path.name, "Failed to decode file with common encodings"
                )

            if not content.strip():
                raise DocumentLoadError(file_path.name, "File is empty")

            # Create metadata
            metadata = self._create_metadata(file_path)
            metadata.update(
                {
                    "loader": "TXTLoader",
                    "encoding": encoding,
                }
            )

            self.logger.info(f"Successfully loaded TXT: {file_path.name}")

            return Document(content=content, metadata=metadata)

        except DocumentLoadError:
            raise
        except Exception as e:
            raise DocumentLoadError(file_path.name, str(e))


class CSVLoader(BaseLoader):
    """Loader for CSV files."""

    def load(self, file_path: Path) -> Document:
        """Load CSV document."""
        try:
            self.logger.info(f"Loading CSV: {file_path.name}")

            # Read CSV with pandas
            df = pd.read_csv(file_path)

            if df.empty:
                raise DocumentLoadError(file_path.name, "CSV file is empty")

            # Convert DataFrame to readable text format
            content_parts = []

            # Add summary
            content_parts.append(f"CSV Data Summary:")
            content_parts.append(f"Total Rows: {len(df)}")
            content_parts.append(f"Total Columns: {len(df.columns)}")
            content_parts.append(f"Columns: {', '.join(df.columns)}")
            content_parts.append("\n")

            # Add data as formatted table
            content_parts.append("Data:")
            content_parts.append(df.to_string(index=False))

            content = "\n".join(content_parts)

            # Create metadata
            metadata = self._create_metadata(file_path)
            metadata.update(
                {
                    "total_rows": len(df),
                    "total_columns": len(df.columns),
                    "columns": list(df.columns),
                    "loader": "CSVLoader",
                }
            )

            self.logger.info(
                f"Successfully loaded CSV: {file_path.name} ({len(df)} rows, {len(df.columns)} columns)"
            )

            return Document(content=content, metadata=metadata)

        except Exception as e:
            raise DocumentLoadError(file_path.name, str(e))


class DOCXLoader(BaseLoader):
    """Loader for Microsoft Word documents."""

    def load(self, file_path: Path) -> Document:
        """Load DOCX document."""
        try:
            self.logger.info(f"Loading DOCX: {file_path.name}")

            try:
                from docx import Document as DocxDocument
            except ImportError:
                raise DocumentLoadError(
                    file_path.name,
                    "python-docx package not installed. Install with: pip install python-docx",
                )

            doc = DocxDocument(file_path)

            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)

            if not content.strip():
                raise DocumentLoadError(
                    file_path.name, "No text content extracted from DOCX"
                )

            # Create metadata
            metadata = self._create_metadata(file_path)
            metadata.update(
                {
                    "total_paragraphs": len(paragraphs),
                    "loader": "DOCXLoader",
                }
            )

            # Try to extract document core properties
            try:
                core_props = doc.core_properties
                metadata.update(
                    {
                        "title": core_props.title or "",
                        "author": core_props.author or "",
                        "subject": core_props.subject or "",
                    }
                )
            except Exception:
                pass

            self.logger.info(f"Successfully loaded DOCX: {file_path.name}")

            return Document(content=content, metadata=metadata)

        except DocumentLoadError:
            raise
        except Exception as e:
            raise DocumentLoadError(file_path.name, str(e))


class XLSXLoader(BaseLoader):
    """Loader for Excel files."""

    def load(self, file_path: Path) -> Document:
        """Load XLSX document."""
        try:
            self.logger.info(f"Loading XLSX: {file_path.name}")

            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names

            content_parts = []
            content_parts.append(f"Excel File with {len(sheet_names)} sheet(s)")
            content_parts.append("\n")

            all_sheets_data = {}

            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                content_parts.append(f"## Sheet: {sheet_name}")
                content_parts.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
                content_parts.append(f"Columns: {', '.join(df.columns)}")
                content_parts.append("\n")
                content_parts.append(df.to_string(index=False))
                content_parts.append("\n" + "=" * 50 + "\n")

                all_sheets_data[sheet_name] = {
                    "rows": len(df),
                    "columns": list(df.columns),
                }

            content = "\n".join(content_parts)

            # Create metadata
            metadata = self._create_metadata(file_path)
            metadata.update(
                {
                    "total_sheets": len(sheet_names),
                    "sheet_names": sheet_names,
                    "sheets_data": all_sheets_data,
                    "loader": "XLSXLoader",
                }
            )

            self.logger.info(
                f"Successfully loaded XLSX: {file_path.name} ({len(sheet_names)} sheets)"
            )

            return Document(content=content, metadata=metadata)

        except Exception as e:
            raise DocumentLoadError(file_path.name, str(e))


class LlamaParseLoader(BaseLoader):
    """
    Advanced PDF loader using LlamaParse for complex document extraction.
    
    Features:
    - Preserves table structure as markdown
    - Describes charts and images
    - Handles multi-column layouts
    - Removes headers/footers noise
    """
    
    _parser = None  # Class-level parser cache
    
    def __init__(self):
        """Initialize LlamaParse loader."""
        from config.settings import settings
        
        if not settings.llama_cloud_api_key:
            raise ValueError("LLAMA_CLOUD_API_KEY not configured")
        
        # Lazy import to avoid issues when llama-parse not installed
        try:
            from llama_parse import LlamaParse
        except ImportError:
            raise ImportError(
                "llama-parse package not installed. Install with: pip install llama-parse"
            )
        
        # Create parser with optimal settings for complex PDFs
        if LlamaParseLoader._parser is None:
            LlamaParseLoader._parser = LlamaParse(
                api_key=settings.llama_cloud_api_key,
                result_type=settings.llamaparse_result_type,
                parsing_instruction=(
                    "Extract all content from this document. "
                    "Format tables as markdown tables with headers. "
                    "For charts and graphs, provide a detailed description of what they show. "
                    "Preserve section headers and document structure. "
                    "Remove repeated headers and footers."
                ),
                verbose=False,
            )
        
        self.parser = LlamaParseLoader._parser
    
    def load(self, file_path: Path) -> Document:
        """Load PDF document using LlamaParse."""
        import asyncio
        
        try:
            self.logger.info(f"Loading PDF with LlamaParse: {file_path.name}")
            
            # Run async parse in sync context
            try:
                loop = asyncio.get_event_loop()
                if loop.is_running():
                    # We're in an async context (like Streamlit)
                    import nest_asyncio
                    nest_asyncio.apply()
                    documents = loop.run_until_complete(
                        self.parser.aload_data(str(file_path))
                    )
                else:
                    documents = loop.run_until_complete(
                        self.parser.aload_data(str(file_path))
                    )
            except RuntimeError:
                # No event loop exists, create one
                documents = asyncio.run(self.parser.aload_data(str(file_path)))
            
            if not documents:
                raise DocumentLoadError(
                    file_path.name, "LlamaParse returned no content"
                )
            
            # Combine all parsed documents
            content_parts = []
            for i, doc in enumerate(documents):
                if hasattr(doc, 'text') and doc.text.strip():
                    content_parts.append(doc.text)
            
            content = "\n\n".join(content_parts)
            
            if not content.strip():
                raise DocumentLoadError(
                    file_path.name, "No text content extracted from PDF"
                )
            
            # Create metadata
            metadata = self._create_metadata(file_path)
            metadata.update({
                "total_pages": len(documents),
                "loader": "LlamaParseLoader",
                "parse_method": "llamaparse",
                "result_type": "markdown",
            })
            
            self.logger.info(
                f"Successfully loaded PDF with LlamaParse: {file_path.name} ({len(documents)} pages)"
            )
            
            return Document(content=content, metadata=metadata)
            
        except Exception as e:
            self.logger.error(f"LlamaParse failed for {file_path.name}: {e}")
            raise DocumentLoadError(file_path.name, str(e))


class DocumentLoaderFactory:
    """Factory for creating appropriate document loaders."""

    _loaders = {
        "pdf": PDFLoader,
        "txt": TXTLoader,
        "csv": CSVLoader,
        "docx": DOCXLoader,
        "xlsx": XLSXLoader,
        "xls": XLSXLoader,
    }

    @classmethod
    def get_loader(cls, file_path: Path) -> BaseLoader:
        """
        Get appropriate loader for file type.

        Args:
            file_path: Path to file

        Returns:
            Loader instance

        Raises:
            UnsupportedFileTypeError: If file type not supported
        """
        from config.settings import settings
        
        extension = get_file_extension(file_path.name)
        
        # Use LlamaParse for PDFs if enabled and API key available
        if extension == "pdf" and settings.enable_llamaparse and settings.llama_cloud_api_key:
            try:
                return LlamaParseLoader()
            except (ImportError, ValueError) as e:
                # Fall back to standard PDFLoader
                import logging
                logging.getLogger(__name__).warning(
                    f"LlamaParse unavailable, falling back to PDFLoader: {e}"
                )

        loader_class = cls._loaders.get(extension)
        if not loader_class:
            raise UnsupportedFileTypeError(file_path.name, extension)

        return loader_class()

    @classmethod
    def load_document(cls, file_path: Path) -> Document:
        """
        Load document using appropriate loader.

        Args:
            file_path: Path to file

        Returns:
            Loaded document

        Raises:
            DocumentLoadError: If loading fails
            UnsupportedFileTypeError: If file type not supported
        """
        loader = cls.get_loader(file_path)
        return loader.load(file_path)


__all__ = [
    "Document",
    "BaseLoader",
    "PDFLoader",
    "LlamaParseLoader",
    "TXTLoader",
    "CSVLoader",
    "DOCXLoader",
    "XLSXLoader",
    "DocumentLoaderFactory",
]
